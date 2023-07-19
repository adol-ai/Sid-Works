import os
import csv
import docx
import time
import numpy
import webbrowser
from PIL import Image
import mysql.connector
from datetime import datetime

exec_Time=datetime.now().strftime("%r").replace(":",";")
exec_Date=datetime.now().strftime("%d%h%G")
doc=docx.Document()
sec=doc.sections
for x in sec:
    x.top_margin=docx.shared.Cm(1.27)
    x.bottom_margin=docx.shared.Cm(1.27)
    x.right_margin=docx.shared.Cm(1.27)
    x.left_margin=docx.shared.Cm(1.27)

s_=' '
gs="@$#&%*+=:^;~-.{}".format(s_)
os.chdir("I:\\Sid\\Python Files\\Internship\\Vimming\\Images")
for i in os.walk(os.getcwd()):
    dir_=i[0]
    images=i[2]
imgc=0
pages=[]
T_data=[]
csvData=[]
csvData_=[]
for pic in images:
    imgc+=1
    print("\n\n\n-->",pic)
    loc=dir_+"\\"+pic
    img=Image.open(loc).convert('L')
    ar=numpy.array(img)
    pic_w,pic_h=ar.shape
    avg_data=numpy.average(ar.reshape(pic_w,pic_h))
    csvSet=[None, None, None, None]
    types=[80,150,250,pic_w]
    imgSet=[]
    T_type=0
    csv_c=-1
    cols_w=0
    Doc_t=0
    for cols in types:
        T_Col1=time.time()
        if cols>pic_w:
            cols=pic_w
            cols_w+=1
            if csvSet[len(csvSet)-1]!=T_type:
                csvSet[len(csvSet)-1]=pic_w
                print("Corrected Width:",cols)
            types.pop()
        if cols_w<=1:
            scale=0.47
            width=pic_w/cols
            height=width/scale
            rows=int(pic_h/height)
            imgList=[]
            for i in range(rows):
                y=int(i*height)
                y_=int((i+1)*height)
                imgList.append("")
                Val=""
                for j in range(cols):
                    x=int(j*width)
                    x_=int((j+1)*width)
                    img_=img.crop((x,y,x_,y_))
                    ar=numpy.array(img_)
                    w,h=ar.shape
                    avg_data=numpy.average(ar.reshape(w,h))
                    gsVal=gs[int((avg_data*14)/255)]
                    Val+=gsVal
                Val='\n'+Val
                if len(set(list(Val)))==2:
                    pass
                else:
                    imgList[i]=Val
            imgSet+=imgList+['<br><br><br><br><br>']
            imgData=''.join(imgList)
            size_dict={50:7.5,100:5.5,150:4.5,200:3.5,250:3,\
            300:2.5,350:2,400:1.5,450:1} 
            for y in size_dict.keys():
                if cols/y<=1:
                    px=size_dict[y]
                    break
                else:
                    px=size_dict[450]
            T_Col2=time.time()
            chrC=rows*cols
            T_type=T_Col2-T_Col1
            print("Resolution Type: ",cols," Done!")
            print(">No. Of Characters: ", chrC)
            print("Execution time: ",T_type,"Secs\n")
            if cols!=pic_w:
                csv_c+=1
                csvSet[csv_c]=T_type
            else:
                csvSet[len(csvSet)-1]=T_type
            T_doc1=time.time()
            doc.add_heading("--->"+pic+" CharacterCount: "+str(chrC))
            doc.add_paragraph("\n\n\n\n\n\n\n\n")
            para=doc.add_paragraph()
            p=para.add_run(imgData)
            p.font.size=docx.shared.Pt(px)
            p.font.name='Courier New'
            para.paragraph_format.line_spacing=docx.shared.Cm(0)
            doc.add_page_break()
            T_doc2=time.time()
            T_doc=(T_doc2-T_doc1)
            Doc_t+=T_doc
    print("\nWord Doc ",pic," Written!")
    print("Time Consumed: ",Doc_t,"Secs\n")
    
    T_page1=time.time()
    pages.append('imagePage%s.html'%str(imgc))
    fhtml=open('I:\\Sid\\Python Files\\Internship\\Vimming\\WebPages\\imagePage%s.html'%str(imgc),'w')
    ip_set=["<html>","<meta name='viewport' content='device-width, initial-scale=1.0'>",\
    "<title>",pic,"</title>","<link rel='icon' type='image/x-icon' href='%s'>"%(loc),\
    "<body>","<style>","*{background-color:grey}",\
    "pre{text-align:center;font-family:Courier New;font-weight:700;font-size: %dpx}"%(px+2),
    "</style>","<pre>"]
    ip_set_=["</pre>","</body>","</html>"]
    complete_ip=ip_set+imgSet+ip_set_
    for io in complete_ip:
        fhtml.write(io)
    fhtml.close()
    T_page2=time.time()
    T_page=T_page2-T_page1
    print("WebPage Created")
    print("Time for Creation: ",T_page,"Secs\n")
    
    csvSet=[pic, round(os.path.getsize(loc)/1024, 2), pic_w]+csvSet
    csvData.append(csvSet+[T_doc, T_page])
    csvData_.append([exec_Date, exec_Time]+csvSet+[T_doc, T_page])
    
doc.save("I:\\Sid\\Python Files\\Internship\\Vimming\\DaWm.docx")


with open("I:\\Sid\\Python Files\\Internship\\Vimming\\CsvFiles\\CsvFile_{}.csv".format(exec_Date+\
"__"+exec_Time),'w') as csvf:
    w=csv.writer(csvf)
    h=["ImgFileName", "FileSize(KB)", "OriginalType_Cols", "80Cols_Type(Secs)", "150Cols_Type(Secs)",\
    "250Cols_Type(Secs)", "Original_Type(Secs)", "Writing_Doc(Secs)", "Writing_WebPage(Secs)"]
    w.writerow(h)
    w.writerows(csvData)

with open("I:\\Sid\\Python Files\\Internship\\Vimming\\CsvFiles\\MainSet\\CompleteCsvData.csv",'a') as csvf:
    w=csv.writer(csvf)
    w.writerows(csvData_)

con=mysql.connector.connect(host="localhost", user="root", passwd="12345", database="Internship")
cur=con.cursor()
try:
    cur.execute("drop table AnalysisDataSet")
except:
    pass
cur.execute("create table AnalysisDataSet(Date varchar(10) not null, Time varchar(15) not null,\
            ImageFileName varchar(200) not null, FileSize_KB float not null, OriginalType_Cols int not null,\
            80Cols_Type_Secs varchar(30), 150Cols_Type_Secs varchar(30), 250Cols_Type_Secs varchar(30),\
            Original_Type_Secs float, Writing_Doc_Secs varchar(30), Writing_WebPage_Secs varchar(30))")
with open("I:\\Sid\\Python Files\\Internship\\Vimming\\CsvFiles\\MainSet\\CompleteCsvData.csv",'r') as csvf:
    r=csv.reader(csvf)
    for i in r:
        if i!=[]:
            for j in i:
                if j=='':
                    i[i.index(j)]="NULL"
            print(tuple(i))
            print(len(i))
            cur.execute("insert into AnalysisDataSet values{}".format(tuple(i)))
con.commit()

for page in pages:
    loc="I:\\Sid\\Python Files\\Internship\\Vimming\\WebPages\\"+page
    webbrowser.open(loc)
    
