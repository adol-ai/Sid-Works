import os
import docx
import numpy
import pprint
from PIL import Image

doc=docx.Document()
sec=doc.sections
for x in sec:
    x.top_margin=docx.shared.Cm(1.27)
    x.bottom_margin=docx.shared.Cm(1.27)
    x.right_margin=docx.shared.Cm(1.27)
    x.left_margin=docx.shared.Cm(1.27)

gs="@$#&%*+=;-~."

img=Image.open("D:\\Python Files\\Internship\\Vimming\\Images\\ChaD.jpeg").convert('L')
ar=numpy.array(img)
pic_w,pic_h=ar.shape
avg_data=numpy.average(ar.reshape(pic_w,pic_h))

cols=500
if cols>pic_w:
    cols=pic_w
    print("Corrected Width:",cols)
    #types.pop()
scale=0.47
width=pic_w/cols
height=width/scale
rows=int(pic_h/height)

imgList=[]
for i in range(rows):
    y=int(i*height)
    y_=int((i+1)*height)
    imgList.append("")
    Val=""
    for j in range(cols):
        x=int(j*width)
        x_=int((j+1)*width)
        img_=img.crop((x,y,x_,y_))
        ar=numpy.array(img_)
        w,h=ar.shape
        avg_data=numpy.average(ar.reshape(w,h))
        gsVal=gs[int((avg_data*11)/255)]
        Val+=gsVal
    Val='\n'+Val
    if len(set(list(Val)))==2:
        pass
    else:
        imgList[i]=Val

#pprint.pprint(imgList)

size_dict={50:7.5,100:5.5,150:4.5,200:3.5,250:3,300:2.5,350:2,400:1.5,450:1}
for y in size_dict.keys():
    if cols/y<=1:
        px=size_dict[y]
        break
    else:
        px=size_dict[450]
        
imgData=''.join(imgList)
#doc.add_heading("--->"+pic)
doc.add_paragraph("\n\n\n\n\n\n\n\n")
para=doc.add_paragraph()
p=para.add_run(imgData)
p.font.size=docx.shared.Pt(px)
p.font.name='Courier New'
para.paragraph_format.line_spacing=docx.shared.Cm(0)
print("Col Type: ",cols," Done!\n")
doc.add_page_break()


fhtml=open('page.html','w')
ip_set=["<html>","<meta name='viewport' content='device-width, initial-scale=1.0'>","<body>","<style>","*{background-color:grey}","pre{text-align:center;font-family:Courier New;font-weight:700;font-size: %dpx}"%(px+2),"</style>","<pre>"]
ip_set_=["</pre>","</body>","</html>"]
complete_ip=ip_set+imgList+ip_set_
for io in complete_ip:
    fhtml.write(io)
fhtml.close()

doc.save("D:\\Python Files\\Internship\\Vimming\\DaWm.docx")

        


